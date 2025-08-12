from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt

def _insert_paragraph_after(paragraph, text="", style=None):
    """
    Insert a new paragraph directly after the given paragraph.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    run = new_para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    return new_para

def add_docx_comment(doc, para_index, comment_text):
    """
    Add an inline comment paragraph after the specified paragraph index.
    Replaces unsupported doc.comments.add().
    """
    try:
        if para_index < 0 or para_index >= len(doc.paragraphs):
            para = doc.paragraphs[-1]
        else:
            para = doc.paragraphs[para_index]
        _insert_paragraph_after(para, f"COMMENT (ADGM Compliance Bot): {comment_text}")
    except Exception as e:
        print(f"Error adding inline comment to paragraph {para_index}: {e}")